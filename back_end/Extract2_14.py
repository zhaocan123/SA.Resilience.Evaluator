"""
在para.name不在标题列表里面的情况下，则必不会是标题
"""
import os
import shutil
import tkinter as tk
import zipfile
from tkinter import filedialog
import pandas as pd
import win32com.client as win32
import xml.dom.minidom
from win32com.client import constants  # visio信息提取
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from xml.dom.minidom import parse
import networkx as nx
import matplotlib.pyplot as plt
import cn2an  # 中文字符转阿拉伯数字
from pylab import *
from docx import Document
from lxml import etree as ET
from xml.dom.minidom import parseString
import re

import docx
from docx.text.paragraph import Paragraph
from docx.oxml.ns import nsmap, qn
from docx.image.image import Image
from docx.parts.image import ImagePart
from docx.oxml.shape import CT_Picture

"""Node类，存储包括标题、文本、表格、图片4种信息"""


class Node:
    def __init__(self, id, name='', text='', father_id=[], father_text=[], son_id=[], son_text=[], level=0, prefix='', table_info={}):
        # 下面为Node对象增加n个实例变量
        self.id = id  # 自己的ID
        self.name = name  # 节点的名称
        self.text = text  # 存储的文本信息
        self.father_id = father_id  # 父亲们的ID
        self.father_text = father_text  # 父亲们的名字
        self.son_id = son_id  # 儿子们的ID
        self.son_text = son_text  # 儿子们的名字
        self.level = level  # 所在层
        self.prefix = prefix  # 前缀
        self.table_info = table_info  # 表格内容（保存表格的字典格式）

    def node_info_print(self):  # 打印各节点信息
        print('')
        print("\033[1;47m\033[1;31m" + '节点id：' + str(self.id) + "\033[0m")
        print('节点title：' + "\033[1;" + str(31 + self.level) + "m" + str(self.name) + "\033[0m")
        print('节点文本信息：' + "\033[1;" + str(31 + self.level) + "m" + str(self.text) + "\033[0m")
        print('所在层：' + "\033[1;" + str(31 + self.level) + "m" + str(self.level) + "\033[0m")
        print('子前缀：' + str(self.prefix))
        print('父节点id：' + str(self.father_id))
        print('父节点文本信息：' + str(self.father_text))
        print('子节点id：' + str(self.son_id))
        print('子节点文本信息：' + str(self.son_text))


# 全局唯一标识
unique_id = 1

'''
打开选择文件夹对话框，读取docx文件，删除已经存在的zip以及解压文件夹
'''


def Get_filepath():
    root = tk.Tk()
    root.withdraw()
    Filepath = filedialog.askopenfilename()  # 获得选择好的文件
    if Filepath == '':
        exit()
    else:
        Filepath = Filepath.replace("\\", "/")  # 路径字符串斜杠替换
    (Folderpath, Filename) = os.path.split(Filepath)  # 分离文件夹路径和文件名.后缀

    # 为了后续copy文件改为zip并解压，删除已经存在的同名文件
    (souce_name, souce_suffix) = os.path.splitext(Filename)  # 分离文件名和后缀
    if souce_suffix != '.docx':
        exit('请选择docx文档')
    extract_path = Folderpath + '/' + souce_name + '_copy'
    zip_path = Folderpath + '/' + souce_name + '_copy.zip'
    Delete_Copy(extract_path, zip_path)  # 删除该路径下的文件
    print('docx文件路径:', Filepath)

    return Filepath


"""
删除文件夹和zip压缩包
"""


def Delete_Copy(extract_path, zip_path):
    # 删除解压文件夹
    if os.path.exists(extract_path):  # 如果文件存在
        # 删除文件夹，可使用以下两种方法。
        shutil.rmtree(extract_path)  # 执行删除文件夹
        print('\n已删除文件夹:' + extract_path)
    else:
        print('no such file:' + extract_path)  # 则返回文件不存在

    # 删除zip压缩包
    if os.path.exists(zip_path):  # 如果文件存在
        # 删除文件夹，可使用以下两种方法。
        os.remove(zip_path)  # 删除zip文件
        print('已删除压缩包:' + zip_path)
    else:
        print('no such file:' + zip_path)  # 则返回文件不存在


"""
从docx文件中提取xml文档，输出解压前zip文件以及解压后文件夹路径
"""


def Get_xmlmkdir(Filepath):
    # 分离文件夹路径、文件名、后缀
    (Folderpath, Filename) = os.path.split(Filepath)  # 分离文件夹路径和文件名.后缀
    (souce_name, souce_suffix) = os.path.splitext(Filename)  # 分离文件名和后缀

    # 克隆一份docx文件，用于改后缀并解压
    os.chdir(Folderpath)
    app = win32.Dispatch('Word.Application')
    source_doc = app.Documents.Open(Filepath)  # 源docx文件路径
    source_doc.Content.Copy()
    copy_name = 'copy'
    copy_path = Folderpath + '/' + souce_name + '_' + copy_name + '.docx'  # copy的docx文件路径
    source_doc.SaveAs(copy_path)  # 另存
    source_doc.Close()

    # 该后缀改成zip文件
    zip_path = Folderpath + '/' + souce_name + '_' + copy_name + '.zip'  # zip路径
    os.rename(copy_path, zip_path)  # 重命名为zip文件

    # 创建新文件夹并把zip内容解压进来
    extract_path = Folderpath + '/' + souce_name + '_' + copy_name
    os.mkdir(extract_path)  # 创建目录
    f = zipfile.ZipFile(zip_path, 'r')  # 进行解压
    for file in f.namelist():
        f.extract(file, extract_path)

    # 输出解压信息
    print('\n源docx文件路径：' + Filepath)
    print('克隆docx文件路径：' + copy_path)
    print('zip文件路径：' + zip_path)
    print('解压文件夹：' + extract_path)
    print('')

    return extract_path, zip_path, Folderpath, souce_name


"""
打印列表各节点
"""


def Print_TreeList(TreeList):
    for each_node in TreeList:  # 遍历节点
        each_node.node_info_print()  # 打印各节点信息


"""
将结点信息存入dataframe中
"""


def TreeList_to_df(TreeList):
    dfname = ['节点id', 'Title', '文本信息', '层级', '父节点id', '父节点文本信息', '子节点id', '子节点文本信息']
    df = pd.DataFrame(columns=dfname)  # 初始化dataframe
    index = 0
    for node in TreeList:
        if node.name == 'Root':  # 不存根节点
            continue
        tmp_list = []
        tmp_list.append(str(node.id))
        tmp_list.append(str(node.name))
        tmp_list.append(str(node.text))
        tmp_list.append(str(node.level))
        tmp_list.append(str(node.father_id))
        tmp_list.append(str(node.father_text))
        if node.son_id == []:
            tmp_list.append('')
        else:
            tmp_list.append(str(node.son_id))
        if node.son_text == []:
            tmp_list.append('')
        else:
            tmp_list.append(str(node.son_text))
        df.loc[index] = tmp_list
        index = index + 1
    return df


"""
将dataframe数据存入excel，居中，自动调整列宽，无合并单元格
"""


def to_excel_nomerger(df, Filepath):
    wb = Workbook()  # 创建一个workbook
    ws = wb.active  # 获取当前workbook的第一个worksheet，默认的索引值是0，它是可以改变的

    # 将每行数据写入ws中
    for row in dataframe_to_rows(df, index=False, header=True):
        ws.append(row)

    # 遍历ws.cell，设置居中
    for cell_i in range(df.shape[0] + 1):
        for cell_j in range(df.shape[1]):
            ws.cell(row=cell_i + 1, column=cell_j + 1).alignment = Alignment(horizontal='center', vertical='center')

    # 自动调整列宽
    for col in df.columns:
        index = list(df.columns).index(col)  # 列序号
        letter = get_column_letter(index + 1)  # 列字母
        collen = df[col].apply(lambda x: len(str(x).encode())).max()  # 获取这一列长度的最大值 当然也可以用min获取最小值 mean获取平均值
        if collen > 150:
            collen = 150  # 限宽
        elif collen < 20:
            collen = 20
        ws.column_dimensions[letter].width = collen

    # 存入excel
    (Folderpath, Filename) = os.path.split(Filepath)  # 分离文件夹路径和文件名.后缀
    (souce_name, souce_suffix) = os.path.splitext(Filename)  # 分离文件名和后缀
    excel_path = Folderpath + '/' + souce_name + '_info_extraction.xlsx'
    if os.path.exists(excel_path):  # 如果文件存在
        # 删除文件夹，可使用以下两种方法。
        os.remove(excel_path)  # 删除zip文件
        print('\n已删除旧excel')
    wb.save(excel_path)
    # 保存writer中的数据至excel
    print('已将节点信息导入excel  ' + excel_path)


"""
传入树列表生成一个图
"""


def Structural_network(TreeList):
    fig, ax = plt.subplots()
    G = nx.DiGraph()  # 无多重边有向图
    max_level = 0
    mycolor = ['#ff0000', 'dodgerblue', 'gold', 'limegreen', 'c', 'lightpink', 'orange', 'violet']
    all_color = []
    node_info = {}
    all_pos = {}
    for node in TreeList:
        if node.level > max_level:
            max_level = node.level
    for level in range(max_level + 1):
        tmpG = nx.DiGraph()
        if level < len(mycolor):
            current_color = mycolor[level]
        else:
            current_color = 'grey'
        for node in TreeList:
            if node.level == level:
                G.add_node(node.id, name=node.name, id=node.id)
                tmpG.add_node(node.id, name=node.name, id=node.id)
                all_color.append(current_color)
                node_info[node.id] = [node.name, node.text, node.level, node.father_id, node.son_id]
        R = level + pow(1.5, level + 1)  # 设置同心圆半径
        if R > 15:
            R = 15
        pos = nx.circular_layout(tmpG, scale=R)
        if len(pos) == 1:
            for key in pos:
                pos[key] = [R, 0]
        if level == 0:
            all_pos = pos
        else:
            all_pos.update(pos)
    for node in TreeList:
        begin_id = node.id
        all_end_id = node.son_id
        if all_end_id != []:
            for i, end_id in enumerate(all_end_id):
                G.add_edges_from([(begin_id, end_id)])
    nx.draw(G, all_pos, node_size=150, node_color=all_color, alpha=0.5)
    node_labels = nx.get_node_attributes(G, 'id')
    nx.draw_networkx_labels(G, all_pos, labels=node_labels, font_size=6)

    po_annotation = []
    for key in all_pos:
        # 标注点的坐标
        point_x = all_pos[key][0]
        point_y = all_pos[key][1]
        point, = plt.plot(point_x, point_y, '.', c='red', markersize=0.1, alpha=0.1)
        # 标注plt.annotate
        if node_info[key][3] == []:
            father_id = '无'
        else:
            father_id = str(node_info[key][3])
        if node_info[key][4] == []:
            son_id = '无'
        else:
            son_id = str(node_info[key][4])
        display_info = 'title:' + node_info[key][0] + '\ntext:' + str(node_info[key][1]) + '\nlevel:' + str(
            node_info[key][2]) + '\nfather_id:' + father_id + '\nson_id:' + son_id
        annotation = plt.annotate(display_info, xy=(point_x, point_y), size=7,
                                  bbox=dict(boxstyle='round,pad=0.5', fc='yellow', ec='k', lw=1, alpha=0.3))
        # 默认鼠标未指向时不显示标注信息
        annotation.set_visible(False)
        po_annotation.append([point, annotation])

    def on_move(event):
        visibility_changed = False
        for point, annotation in po_annotation:
            should_be_visible = (point.contains(event)[0] == True)
            if should_be_visible != annotation.get_visible():
                visibility_changed = True
                annotation.set_visible(should_be_visible)
        if visibility_changed:
            plt.draw()

    on_move_id = fig.canvas.mpl_connect('motion_notify_event', on_move)
    mpl.rcParams['font.sans-serif'] = ['SimHei']  # ''' 添加中文字体'''
    plt.rcParams['axes.unicode_minus'] = False  # ''' 防止负号无法正常显示  ''''
    plt.show()


def Output_to_txt(TreeList, Folderpath, souce_name):
    print(Folderpath)
    txt_path = Folderpath + '/' + souce_name + '.txt'

    # 删除存在的txt文档
    if os.path.exists(txt_path):  # 如果文件存在
        # 删除文件夹，可使用以下两种方法。
        os.remove(txt_path)  # 删除zip文件
        print('已删除文档:' + txt_path)
    else:
        print('no such file:' + txt_path)  # 则返回文件不存在

    f = open(txt_path, 'w+', encoding='utf-8')
    for node in TreeList:
        if node.level > 0:
            suojin = ''
            i = 1
            while i < node.level:
                suojin = suojin + '    ' + '    '
                i = i + 1
            f.write(suojin + '【' + str(node.level) + ':' + node.name + '|' + node.prefix + '】' + str(node.text) + '\n')
    f.close()


# 遍历所有的节点
def walkData(root_node, level, result_list):
    if (level > 3):
        return
    global unique_id
    temp_list = [root_node.tag.split('}')[1]]
    result_list.append(temp_list)
    unique_id += 1

    # 遍历每个子节点
    children_node = root_node.getchildren()
    if len(children_node) == 0:
        return
    for child in children_node:
        walkData(child, level + 1, result_list)
    return


def getXmlData(file_name):
    level = 1  # 节点的深度从1开始
    result_list = []
    root = ET.parse(file_name).getroot()
    walkData(root, level, result_list)

    return result_list


'''
提取顺序
'''


def get_order(extract_path):
    Parsing_path = extract_path + '/word/document.xml'
    Tag_order = getXmlData(Parsing_path)
    index_p = 0
    index_tbl = 0
    index_sdt = 0
    for Tag in Tag_order:
        if Tag[0] == 'p':  # 文本、图片
            Tag.append(index_p)
            # if index_p == 273:
            #     print(Tag)
            index_p = index_p + 1
            print(Tag)

        elif Tag[0] == 'tbl':  # 表格
            Tag.append(index_tbl)
            index_tbl = index_tbl + 1
            print(Tag)
        elif Tag[0] == 'sdt':  # 目录
            Tag.append(index_sdt)
            index_sdt = index_sdt + 1
            print(Tag)
    return Tag_order


"""
document 为文档对象
paragraph 为内嵌图片的段落对象
"""


def get_vsdx_picture(paragraph: Paragraph, extract_path):
    text = ''
    img_vsdx = paragraph._element.xpath('.//w:object')
    if img_vsdx:
        img_vsdx_xml = paragraph._element.xml
        img_vsdx_xml = ET.fromstring(img_vsdx_xml)
        img_vsdx_dom = parseString(ET.tounicode(img_vsdx_xml))
        img_vsdx_OLEObject = img_vsdx_dom.getElementsByTagName('o:OLEObject')
        if img_vsdx_OLEObject:
            rid = img_vsdx_OLEObject[0].getAttribute('r:id')
            if rid != '':
                partname = paragraph.part.related_parts[rid].partname
                vsdx_path = extract_path + partname
                text = Get_vsdx_from_path(vsdx_path)
    return text


"""
输入图片存放路径获取图片信息
"""


def Get_vsdx_from_path(picture_path):
    (Folderpath, Filename) = os.path.split(picture_path)  # 分离文件夹路径和文件名.后缀
    (souce_name, souce_suffix) = os.path.splitext(Filename)  # 分离文件名和后缀
    appVisio = win32.gencache.EnsureDispatch("Visio.Application")
    if souce_suffix == '.vsdx':
        vdoc = appVisio.Documents.Open(picture_path)  # 打开vsdx文档
        connects = vdoc.Pages(1).Connects
        text = []  # 初始化text列表，存储visio图片中的每一条边信息
        for conn in connects:  # 遍历每一条边
            from_shp = conn.FromSheet
            info = from_shp.Name
            if conn.FromPart == constants.visBegin:
                info += ', ' + '起点'
            elif conn.FromPart == constants.visEnd:
                info += ', ' + '终点'  # visBegin ==9 ,visEnd == 12
            to_shp = conn.ToSheet
            info += ', ' + to_shp.Text  # 连接的矩形名称
            text.append(info)
        vdoc.Close()  # 关闭vsdx文档
    else:
        text = '非vsdx文档无法读取'
    return text


"""
提取表格文本信息存入字典中并返回
"""


def Get_table_info(table):
    lo = {}  # 存储每一行去重后的数据
    for row in range(0, len(table.rows)):
        row_list = []
        for col in range(0, len(table.row_cells(row))):  # 提取row行的全部列数据
            row_list.append(
                table.cell(row, col).text.replace('\n', '').replace(' ', ''))  # 去除字符串中的特殊字符，并添加到临时列表中
        # lo[row] = (sorted(set(row_list), key=row_list.index))  # 在不变顺序的前提下，去除List中的重复项
        lo[row] = row_list

        # 打印出每行的数据观察相关格式特征
        # print(row, ":len(", len(lo[row]), '):', lo[row])
    return lo  # 表格字典


"""
针对没有自动编号的前缀，提取其前缀和文本
"""


def Get_prefix_and_text(S):
    ZH = '零一二三四五六七八九十'
    separator = '.、\/-()（）'
    i = 0
    tmp_run = ''
    all_run = []
    prefix_level = 0  # 前缀的层次

    # 将前缀部分的中文数字转为阿拉伯数字
    while i < len(S):
        if S[i] in ZH:
            while S[i] in ZH:
                tmp_run = tmp_run + S[i]
                i = i + 1
                if i >= len(S):
                    break
            all_run.append(tmp_run)
            tmp_run = ''
        else:
            while S[i] not in ZH:
                tmp_run = tmp_run + S[i]
                i = i + 1
                if i >= len(S):
                    break
            all_run.append(tmp_run)
            tmp_run = ''
    # print(all_run)
    for i in range(len(all_run)):
        if all_run[i][0] in ZH:  # 如果是数字，转换为阿拉伯字符
            all_run[i] = str(cn2an.cn2an(all_run[i], "normal"))
        else:
            if len(all_run[i]) > 1:
                break

    # 拼接各块
    Arabic_S = ''.join(all_run)

    # 提取前缀
    prefix = ''
    prefix_feature = ''
    first_in = True
    text = ''
    for i in range(len(Arabic_S)):
        if not Arabic_S[i].isdigit() and not Arabic_S[i].encode('UTF-8').isalpha() and Arabic_S[i] not in separator:
            text = Arabic_S[i:len(Arabic_S)]
            break
        if (Arabic_S[i].isdigit() or Arabic_S[i].encode('UTF-8').isalpha()) and first_in:
            prefix_level = prefix_level + 1
            prefix_feature = prefix_feature + 'x'
            first_in = False
        if Arabic_S[i] in separator:
            first_in = True
            prefix_feature = prefix_feature + Arabic_S[i]
        prefix = prefix + Arabic_S[i]

    return prefix_feature, text


def Get_all_feature(para, doc, all_title_name):
    p_text = para.text
    try:
        p_style = para.style.style_id
        p_size = para.style.font.size
        p_name = para.style.name
    except AttributeError:
        p_style = 'No style id'
        p_size = None
        p_name = 'No para name'

    # 根据name判断是否属于标题类
    if p_name in all_title_name:  # 被认为的标题的段落
        p_istitle = True
    else:  #
        p_istitle = False

    # 提取前缀特征
    try:
        p_numId = para._element.pPr.numPr.numId.val
        p_ilvl = para._element.pPr.numPr.ilvl.val
    except AttributeError:
        p_numId = None
        p_ilvl = None
    if p_numId is None or p_ilvl is None:  # 没有自动编号前缀
        Prefix_feature, p_text = Get_prefix_and_text(p_text)
    else:  # 有自动编号前缀
        ct_numbering = doc.part.numbering_part._element
        numXML = doc.part.numbering_part.numbering_definitions._numbering
        p_abstractId = None
        for num in ct_numbering.num_lst:
            # 获取numId和abstractNmuId的对应关系
            if num.numId == p_numId:
                p_abstractId = num.abstractNumId.val
                break
        if p_abstractId is not None:
            numXML_xml = ET.fromstring(numXML.xml)
            numXML_dom = parseString(ET.tounicode(numXML_xml))
            numXML_dom_abstractNum = numXML_dom.getElementsByTagName('w:abstractNum')
            for each_abstractNum in numXML_dom_abstractNum:
                if each_abstractNum.getAttribute('w:abstractNumId') == str(p_abstractId):
                    numXML_dom_lvl = each_abstractNum.getElementsByTagName('w:lvl')
                    for each_lvl in numXML_dom_lvl:
                        if each_lvl.getAttribute('w:ilvl') == str(p_ilvl):
                            # lvlText
                            numXML_dom_lvlText = each_lvl.getElementsByTagName('w:lvlText')
                            p_lvlText = numXML_dom_lvlText[0].getAttribute('w:val')
            Prefix_feature = re.sub(r'%+\d+', "x", p_lvlText)  # 提取自动编号前缀的特征
            Prefix_feature = re.sub(r'\d+', "x", Prefix_feature)
        else:
            Prefix_feature = ''

    # 构成段落特征
    para_feature = [Prefix_feature, p_style, p_size, p_istitle, p_name, p_text]
    return para_feature


"""
解析docx文档各类型的信息
"""


def Parsing_docx(Tag_order, Filepath, extract_path):
    # 初始化标题name类型列表
    all_title_name = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6', 'Heading 7',
                      'Heading 8', 'Heading 9', 'Title', 'Subtitle', 'Body Text First Indent']

    # 初始化树的根节点，并构造树节点列表
    ID = 0
    Leaf = Node(id=ID, name='Root', text='根节点', level=0, son_id=[], son_text=[])
    ID = ID + 1
    TreeList = []
    TreeList.append(Leaf)
    all_level_title_feature = []
    not_title_tmp = []  # 非标题临时列表
    level = 0

    # 读取文档进行解析
    doc = Document(Filepath)
    Tag_needed = ['p', 'tbl', 'sdt']  # [段落，表格，目录]
    for Tag in Tag_order:  # 遍历保序列后的节点

        if Tag[0] == Tag_needed[0]:  # 处理w:p
            para = doc.paragraphs[Tag[1]]
            image_text = get_vsdx_picture(para, extract_path)
            if image_text == '':  # 文本或标题，且空文本不进行提取
                if para.text != '' and para.text != '\n':
                    para_feature = Get_all_feature(para, doc, all_title_name)  # 提取段落特征[前缀特征、pStyle、字体大小、是否是标题、name、段落文本]
                    compare_feature = [level] + para_feature  # [层级、前缀特征、pStyle、字体大小、是否是标题、name、段落文本]

                    if compare_feature[-3]:  # name在自定义标题列表里面
                        not_title_tmp = []  # 首先清空非标题临时列表
                        if compare_feature[1] != '':  # 如果有前缀
                            nomatch = True  # 默认列表里没有匹配的前缀
                            for each_feature in all_level_title_feature:  # 查找是否能在列表里找到匹配的前缀
                                if each_feature[1] == compare_feature[1]:  # 如果匹配到前缀
                                    level = each_feature[0]
                                    node_level = level
                                    compare_feature[0] = each_feature[0]
                                    all_level_title_feature.append(compare_feature)
                                    nomatch = False
                                    break
                            if nomatch:  # 没有匹配的前缀，就将标题认为是下一层的
                                level = level + 1
                                node_level = level
                                compare_feature[0] = level
                                all_level_title_feature.append(compare_feature)
                        else:  # 没有任何前缀（自动的和自定义的）
                            nomatch = True  # 默认列表里没有匹配的pStyle和字体大小
                            for each_feature in all_level_title_feature:  # 查找是否能在列表里找到匹配的pStyle和字体大小
                                if each_feature[2] == compare_feature[2] and each_feature[3] == compare_feature[
                                        3]:  # 如果匹配到pStyle和字体大小
                                    level = each_feature[0]
                                    node_level = level
                                    compare_feature[0] = each_feature[0]
                                    all_level_title_feature.append(compare_feature)
                                    nomatch = False
                                    break
                            if nomatch:  # 没有匹配的pStyle和字体大小，就将标题认为是下一层的
                                level = level + 1
                                node_level = level
                                compare_feature[0] = level
                                all_level_title_feature.append(compare_feature)
                    else:  # name不在自定义标题列表里面
                        last_node = TreeList[-1]  # 上一个节点
                        if last_node.name == compare_feature[-2] and last_node.prefix == compare_feature[1]:  # 上一个节点是否与自己相同name
                            node_level = last_node.level
                        else:
                            nomatch = True
                            for not_title_traversed in not_title_tmp:  # 在列表里找到匹配name的非标题节点
                                if not_title_traversed[0] == compare_feature[-2] and not_title_traversed[1] == compare_feature[1]:
                                    node_level = not_title_traversed[2]
                                    nomatch = False
                                    break
                            if nomatch:
                                node_level = last_node.level + 1
                                not_title_tmp.append([compare_feature[-2], compare_feature[1], node_level])

                    # 将信息存入叶子节点 #################
                    TreeList.append(
                        Node(id=ID, name=compare_feature[-2], text=compare_feature[-1], father_id=[],
                             father_text=[],
                             son_id=[],
                             son_text=[], level=node_level, prefix=compare_feature[1]))
                    ID = ID + 1
                    ######################################

            else:  # 图片
                last_node = TreeList[-1]  # 上一个节点
                if last_node.name == 'Visio Picture':
                    node_level = last_node.level
                else:
                    nomatch = True
                    for not_title_traversed in not_title_tmp:  # 在列表里找到匹配name的非标题节点
                        if not_title_traversed[0] == 'Visio Picture':
                            node_level = not_title_traversed[1]
                            nomatch = False
                            break
                    if nomatch:
                        node_level = last_node.level + 1
                        not_title_tmp.append(['Visio Picture', node_level])
                # 将图片信息存入叶子节点
                TreeList.append(
                    Node(id=ID, name='Visio Picture', text=image_text, father_id=[], father_text=[], son_id=[],
                         son_text=[], level=node_level))
                ID = ID + 1
                ######################################

        elif Tag[0] == Tag_needed[1]:  # 处理w:tbl
            table = doc.tables[Tag[1]]
            lo = Get_table_info(table)  # 读取表格文本信息存入字典
            last_node = TreeList[-1]  # 上一个节点
            try:
                table_name = table.style.name
            except AttributeError:
                table_name = 'No table name'
            if last_node.name == table_name:
                node_level = last_node.level
            else:
                nomatch = True
                for not_title_traversed in not_title_tmp:  # 在列表里找到匹配name的非标题节点
                    if not_title_traversed[0] == table_name:
                        node_level = not_title_traversed[1]
                        nomatch = False
                        break
                if nomatch:
                    node_level = last_node.level + 1
                    not_title_tmp.append([table_name, node_level])
            # 将表格信息存入叶子节点
            TreeList.append(
                Node(id=ID, name=table_name, text=str(lo), father_id=[], father_text=[], son_id=[], son_text=[],
                     level=node_level, table_info=lo))
            ID = ID + 1
            ######################################

        elif Tag[0] == Tag_needed[2]:  # 处理w:sdt
            p_xmlstr = doc.element.xml
            p_xml = ET.fromstring(p_xmlstr)  # 转换成lxml结点
            p_dom = parseString(ET.tounicode(p_xml))

            sdt_rfonts = p_dom.getElementsByTagName('w:sdt')
            sdt_content_rfonts = sdt_rfonts[Tag[1]].getElementsByTagName('w:sdtContent')
            wp_rfonts = sdt_content_rfonts[0].getElementsByTagName('w:p')
            for wp_i in range(wp_rfonts.length):  # 遍历sdtContent下的每个wp节点
                whyperlink_rfonts = wp_rfonts[wp_i].getElementsByTagName('w:hyperlink')
                if whyperlink_rfonts == []:
                    wt_rfonts = wp_rfonts[wp_i].getElementsByTagName('w:t')
                    wp_text = ''
                    for wt_i in range(wt_rfonts.length):  # 遍历wp下的每个wt节点
                        wt_text = ''
                        for wt_child_i in range(wt_rfonts[wt_i].childNodes.length):
                            wt_text = wt_text + wt_rfonts[wt_i].childNodes[wt_child_i].data
                        wp_text = wp_text + wt_text
                    if wp_text != '':
                        if level == 0:
                            content_level = 1
                        else:
                            content_level = level
                        # 将表格信息存入叶子节点
                        TreeList.append(
                            Node(id=ID, name='Table of Contents', text=wp_text, father_id=[], father_text=[], son_id=[],
                                 son_text=[],
                                 level=content_level))
                        ID = ID + 1
                        ######################################
                else:
                    wt_rfonts = wp_rfonts[wp_i].getElementsByTagName('w:t')
                    wp_text = ''
                    for wt_i in range(wt_rfonts.length):  # 遍历wp下的每个wt节点
                        wt_text = ''
                        for wt_child_i in range(wt_rfonts[wt_i].childNodes.length):
                            wt_text = wt_text + wt_rfonts[wt_i].childNodes[wt_child_i].data
                        wp_text = wp_text + wt_text
                    if wp_text != '':
                        # 将表格信息存入叶子节点
                        TreeList.append(
                            Node(id=ID, name='Content List', text=wp_text, father_id=[], father_text=[], son_id=[],
                                 son_text=[],
                                 level=content_level+1))
                        ID = ID + 1
                        ######################################

    TreeList = Organize_TreeList(TreeList)  # 组织节点中的父子关系
    return TreeList


def Organize_TreeList(TreeList):
    new_TreeList = []
    # 修改儿子与父亲
    for Leaf in TreeList:
        current_level = Leaf.level
        for j in range(len(new_TreeList)):
            father_index = len(new_TreeList) - j - 1
            if new_TreeList[father_index].level <= current_level - 1:  # 找到父节点在列表中所在位置，上一层应该是标题
                Leaf.father_id = new_TreeList[father_index].id
                Leaf.father_text = new_TreeList[father_index].text
                new_TreeList[father_index].son_id.append(Leaf.id)
                new_TreeList[father_index].son_text.append(Leaf.text)
                break
        new_TreeList.append(Leaf)  # 将节点压入列表中
    TreeList = new_TreeList

    return TreeList


if __name__ == '__main__':
    Filepath = Get_filepath()  # 获取docx文件路径
    extract_path, zip_path, Folderpath, souce_name = Get_xmlmkdir(Filepath)  # 将docx的克隆文件解压成含xml文件夹
    Tag_order = get_order(extract_path)
    TreeList = Parsing_docx(Tag_order, Filepath, extract_path)
    Print_TreeList(TreeList)  # 打印树节点信息
    Output_to_txt(TreeList, Folderpath, souce_name)  # 存入txt
    Structural_network(TreeList)  # 生成图
